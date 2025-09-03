import json
import os
import time
from datetime import datetime
import pyautogui
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Crear carpeta para capturas
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
capturas_dir = f"capturas del {timestamp}"
os.makedirs(capturas_dir, exist_ok=True)
extrasleep = 0
# Cargar URLs desde un archivo JSON
print("páginas web")
print("(1) internas")
print("(2) publicadas")
print(":")
fileselected = int(input())
if fileselected == 1:
    json_file = "urls.json"
elif fileselected == 2:
    json_file = "urls-o.json"
else:
    json_file = "ipspub.json"
    extrasleep = fileselected
with open(json_file, "r") as file:
    urls = json.load(file)["urls"]

# Configurar opciones de Edge
edge_options = Options()
edge_options.add_argument("--start-maximized")  # Asegurar pantalla completa

# Obtener la ruta absoluta del archivo msedgedriver.exe en la misma carpeta
driver_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "msedgedriver.exe")

# Iniciar WebDriver de Microsoft Edge
service = Service(driver_path)
driver = webdriver.Edge(service=service, options=edge_options)

if json_file == "urls-o.json":
    index = 23
    tittle = "Monitoreo a páginas web publicadas desde red externa"
    tipoServicio = "página publicada"
else:
    index = 1
    tittle = "Monitoreo a páginas web internas desde red interna"
    tipoServicio = "página interna"

# Crear documento DOCX
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
doc = Document()
paragraph = doc.add_heading(f"{tittle}", level=1)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.runs[0]
run.underline = True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(14)
run.font.name = 'Aptos Display'

# Tomar capturas de pantalla
for i, url in enumerate(urls):
    try:
        driver.get(url)
        time.sleep(4+extrasleep)  # Esperar a que la página cargue completamente

    except Exception as e:
        # Imprimir el error y el enlace defectuoso
        print(f"Error al cargar la página: {url}")
        print(f"Error: {e}")
    
    # Nombre del archivo de captura
    screenshot_path = os.path.join(capturas_dir, f"screenshot_{i+1}_{timestamp}.png")
        
    # Capturar toda la pantalla incluyendo la barra de tareas
    pyautogui.screenshot(screenshot_path)
    print(f"Captura guardada: {screenshot_path}")
    
    # Agregar imagen al documento
    doc.add_picture(screenshot_path, width=Inches(6))  # Ajustar tamaño en el DOCX
    para = doc.add_paragraph()
    run = para.add_run(f"Img. N°{i+index}: {url} ({tipoServicio})")
    run.bold = True
    run.italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

# Cerrar navegador
driver.quit()

# Guardar documento
doc_path = f"capturas_pantallas_{timestamp}.docx"
doc.save(doc_path)
print(f"Documento generado: {doc_path}")
